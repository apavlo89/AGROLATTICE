"""Persistent-evidence publication helpers for AGROLATTICE.

Release 11.15 keeps the legacy public API but modernises the scientific report
builder: richer manuscript metadata, broader figure types, report auditing,
privacy-aware export, method/citation provenance and reproducibility packages.
"""
from __future__ import annotations

import hashlib
import html
import io
import json
import os
import re
import time
import uuid
import zipfile
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Mapping, Sequence

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

MODULE_VERSION = "2.0.0"


class PublicationBuilderError(RuntimeError):
    pass


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def slugify(value: str, fallback: str = "study") -> str:
    text = re.sub(r"[^A-Za-z0-9]+", "-", str(value or "").strip()).strip("-").lower()
    return text[:90] or fallback


def json_safe(value: Any) -> Any:
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    if isinstance(value, (datetime, pd.Timestamp, Path)):
        return str(value)
    if isinstance(value, pd.DataFrame):
        return {"__type__": "dataframe", "columns": [str(c) for c in value.columns], "records": [json_safe(r) for r in value.to_dict("records")]}
    if isinstance(value, pd.Series):
        return {"__type__": "series", "name": str(value.name), "values": json_safe(value.to_dict())}
    if isinstance(value, Mapping):
        return {str(k): json_safe(v) for k, v in value.items()}
    if isinstance(value, (list, tuple, set)):
        return [json_safe(v) for v in value]
    if hasattr(value, "item"):
        try:
            return value.item()
        except Exception:
            pass
    return str(value)


def dataframe_from_json_safe(value: Any) -> pd.DataFrame | None:
    if isinstance(value, pd.DataFrame):
        return value.copy()
    if isinstance(value, Mapping) and value.get("__type__") == "dataframe":
        return pd.DataFrame(value.get("records", []), columns=value.get("columns"))
    return None


REPORT_PRESETS = {
    "Experiment report": ["introduction", "methods", "results", "discussion", "limitations", "conclusion"],
    "Maize Synchrony report": ["introduction", "methods", "results", "discussion", "limitations", "conclusion"],
    "Persistent Twin season report": ["summary", "methods", "results", "limitations", "conclusion"],
    "Field season report": ["summary", "methods", "results", "limitations", "conclusion"],
    "Climate & EO report": ["summary", "methods", "results", "discussion", "limitations", "conclusion"],
    "Crop decision report": ["summary", "methods", "results", "limitations", "conclusion"],
    "Model validation report": ["summary", "methods", "results", "discussion", "limitations", "conclusion"],
    "Recommendation/outcome audit": ["summary", "methods", "results", "limitations", "conclusion"],
    "G×E×M analysis report": ["introduction", "methods", "results", "discussion", "limitations", "conclusion"],
    "Full scientific manuscript": ["introduction", "methods", "results", "discussion", "limitations", "conclusion"],
    "Supplementary/reproducibility package": ["methods", "results", "limitations"],
}

MANUSCRIPT_PRESETS = {
    "Generic IMRaD": "Introduction → Materials & Methods → Results → Discussion → Limitations → Conclusions",
    "Short communication": "Short introduction → Methods → Results & Discussion → Conclusions",
    "Methods paper": "Background → Method → Implementation → Validation → Limitations",
    "Dataset paper": "Context → Data collection → Data records → Technical validation → Usage notes",
    "Technical report": "Executive summary → Methods → Findings → Limitations → Recommendations",
    "Thesis chapter": "Introduction → Methods → Results → Discussion → Conclusions",
}

# "relationship" describes how AGROLATTICE uses the method, not a claim of exact reproduction.
METHODS_REGISTRY: dict[str, dict[str, Any]] = {
    "NASA-derived agroclimate data": {
        "text": "AGROLATTICE used the installed country-scoped NASA-derived agroclimate dataset and/or field-linked NASA POWER retrievals with explicit variable units, temporal coverage and spatial provenance. Gridded environmental values were not represented as local weather-station measurements.",
        "references": ["NASA POWER documentation and the exact AGROLATTICE dataset manifest used in the analysis."],
        "relationship": "Native AGROLATTICE data integration",
    },
    "Daily weather and phenology": {
        "text": "Daily environmental records were used to calculate thermal time and stage-specific exposure variables. Observed, retrieved, derived and modelled phenology were kept distinct in the evidence record.",
        "references": ["NASA POWER Daily Point API documentation where NASA POWER was used."],
        "relationship": "Native AGROLATTICE workflow",
    },
    "Mechanistic Maize Twin": {
        "text": "Maize development used AGROLATTICE's independent mechanistic implementation informed by Laurent et al. (2025): emergence after 30.6 GDD, exponential leaf appearance parameterised by coblf and tln, ear-growth onset at 0.67×tln, female silking linked to ebR1 and male anthesis after final-leaf expansion. Publication priors were kept separate from local measurements and approved calibrations.",
        "references": ["Laurent et al. (2025), Crop Science 65, DOI 10.1002/csc2.21453."],
        "relationship": "Independent paper-derived adaptation; proprietary source data and original C++ Bayesian sampler unavailable",
    },
    "Root-zone soil-water balance": {
        "text": "Daily root-zone depletion was simulated using transparent FAO-56-style water-balance concepts including total and readily available water, crop coefficients, water-stress coefficient, rainfall, irrigation, runoff and deep percolation. Recorded irrigation remained separate from recommendations.",
        "references": ["Allen et al., FAO Irrigation and Drainage Paper 56."],
        "relationship": "Native transparent process model",
    },
    "Sentinel-2 Earth observation": {
        "text": "Sentinel-2 Level-2A imagery was retrieved through public STAC providers for authoritative field polygons. Optical indices were calculated after Scene Classification Layer quality masking, usable-pixel checks and explicit scene/provenance recording.",
        "references": ["Copernicus Sentinel-2 Level-2A product documentation.", "STAC provider documentation for the exact provider recorded in the evidence manifest."],
        "relationship": "Native AGROLATTICE EO processing",
    },
    "Experimental design and randomisation": {
        "text": "Experimental factors, blocks, replication, randomisation seed, allocation manifest and protocol revisions were retained as persistent trial evidence. Analyses were required to respect the declared experimental unit and randomisation structure.",
        "references": [],
        "relationship": "Native AGROLATTICE research workflow",
    },
    "Agricultural grouped validation": {
        "text": "Predictive performance was evaluated using deployment-relevant held-out structures such as field, site, season, trial, genotype, parent-pair, forward-time or spatial blocks rather than relying solely on random row splits. Preprocessing and resampling were confined to training folds.",
        "references": [],
        "relationship": "Native AGROLATTICE scientific-governance framework",
    },
    "Environmental pest-risk modelling": {
        "text": "Environmental pest-risk modelling used transparent engineered weather features, leakage-safe resampling/model selection and model-specific explainability. Risk predictions were kept separate from field confirmation, disease diagnosis and pesticide recommendations.",
        "references": ["Wadhwa & Malik (2024), Computers and Electronics in Agriculture, DOI 10.1016/j.compag.2024.109472."],
        "relationship": "Independent adaptation of published environmental feature/modelling concepts",
    },
    "ALIC pest forecasting": {
        "text": "Where explicitly used, ALIC-style pest forecasting combines feature attention, temporal LSTM representations and interaction convolutional processing of meteorological and historical pest time series. Its applicability requires adequate historical pest observations and walk-forward evaluation.",
        "references": ["Wang & Zhang (2024), Expert Systems with Applications, DOI 10.1016/j.eswa.2024.124137."],
        "relationship": "Paper-derived research method; cite as implemented only when the registered model confirms this architecture",
    },
    "Weakly supervised spatial yield": {
        "text": "Fine-resolution yield estimates were trained under an aggregate-consistency constraint when only coarser outcome labels were available. Fine-scale outputs were labelled weakly supervised unless independently validated at the same spatial support.",
        "references": ["Paudel et al. (2023), Environmental Research Letters, DOI 10.1088/1748-9326/acf50e."],
        "relationship": "Independent adaptation",
    },
    "Adaptive multimodal fusion": {
        "text": "Weather, soil, Earth-observation, management, sensor and phenology modalities were represented separately and combined using held-out reliability-aware fusion. Modality weights were interpreted as predictive reliability signals, not causal effects.",
        "references": [],
        "relationship": "AGROLATTICE multimodal adaptation informed by recent sub-field yield research",
    },
    "Hybrid mechanistic + ML residual learning": {
        "text": "Machine learning was used to model residual error around an existing mechanistic/process prediction. A residual correction was accepted only when it improved held-out performance under the declared agricultural validation design.",
        "references": [],
        "relationship": "Native AGROLATTICE hybrid modelling pattern",
    },
    "Model uncertainty and calibration": {
        "text": "Prediction uncertainty, calibration and applicability were reported using the method actually stored with the model or validation run. Data quality, parameter uncertainty, model disagreement, prediction intervals and out-of-domain status were not collapsed into one generic confidence score.",
        "references": [],
        "relationship": "Native AGROLATTICE model-governance framework",
    },
    "Recommendation causal audit": {
        "text": "Recommendation effectiveness was evaluated separately from predictive accuracy. Observational treatment-effect estimates documented treatment definition, pre-treatment covariates, overlap/positivity diagnostics, grouped cross-fitting and causal assumptions; they were not presented as randomized causal proof.",
        "references": ["Tsoumas et al. (2022), Evaluating Digital Tools for Sustainable Agriculture using Causal Inference, arXiv:2211.03195."],
        "relationship": "Independent adaptation of causal-evaluation principles",
    },
    "Irrigation policy comparison": {
        "text": "Irrigation strategies were compared as explicit scenarios subject to user-specified water, event and operational constraints. Recommended irrigation remained distinct from recorded applied irrigation and automatic external hardware control was not assumed.",
        "references": [],
        "relationship": "Native AGROLATTICE decision workflow",
    },
    "Nutrient response optimisation": {
        "text": "Nutrient-response surfaces and multi-objective alternatives were explored only when treatment-rate variation and independent outcomes were sufficient. Response exploration was separated from operational recommendations and extrapolation beyond observed treatment ranges was flagged.",
        "references": [],
        "relationship": "Native AGROLATTICE decision workflow",
    },
    "AquaCrop-OSPy": {
        "text": "AquaCrop-OSPy was used as an independent Python implementation of AquaCrop-OS. Exact weather, soil, crop and management settings were retained and outputs were not represented as official FAO AquaCrop 7.x runs.",
        "references": ["AquaCrop-OSPy documentation.", "FAO AquaCrop documentation."],
        "relationship": "External process-model backend",
    },
    "DSSAT interoperability": {
        "text": "DSSAT-compatible inputs and execution packages were prepared with explicit executable version, command, working directory, input configuration and output inventory. Successful execution was not treated as proof that cultivar, soil or management parameterisation was valid.",
        "references": ["DSSAT documentation for the installed model version."],
        "relationship": "External process-model interoperability",
    },
    "APSIM interoperability": {
        "text": "APSIM Next Generation simulations were prepared/executed with explicit model version, command and input configuration. Failed runs and configuration limitations were retained as provenance.",
        "references": ["APSIM Next Generation documentation for the installed model version."],
        "relationship": "External process-model interoperability",
    },
    # Legacy aliases retained for older studies.
    "Historical climate data": {"text": "Historical climate records were analysed with explicit spatial/temporal provenance and variable units.", "references": [], "relationship": "Legacy alias"},
    "Soil-water balance": {"text": "Daily root-zone depletion was simulated using transparent crop-water balance assumptions.", "references": ["Allen et al., FAO Irrigation and Drainage Paper 56."], "relationship": "Legacy alias"},
    "Satellite monitoring": {"text": "Sentinel-2 observations were processed for the stored analysis geometry with explicit quality filtering and scene provenance.", "references": [], "relationship": "Legacy alias"},
    "Validation Centre": {"text": "Model performance was evaluated on paired observations and predictions under the declared held-out structure.", "references": [], "relationship": "Legacy alias"},
    "Model ensemble": {"text": "Compatible model predictions were compared/combined using a prespecified rule while retaining model disagreement.", "references": [], "relationship": "Legacy alias"},
    "Water productivity and economics": {"text": "Water productivity and financial outcomes used explicit observed/modelled quantities and user-supplied economic assumptions.", "references": [], "relationship": "Legacy alias"},
}


def new_study_template(*, title: str, project_id: str | None = None, report_type: str = "Full scientific manuscript") -> dict[str, Any]:
    now = utc_now_iso()
    return {
        "schema_version": "2.0.0",
        "study_id": str(uuid.uuid4()),
        "title": title or "Untitled AGROLATTICE research report",
        "short_title": title[:70] if title else "AGROLATTICE research report",
        "project_id": project_id,
        "report_type": report_type,
        "authors": [],
        "affiliations": [],
        "corresponding_author": "",
        "corresponding_orcid": "",
        "journal": "",
        "manuscript_status": "Draft",
        "research_question": "",
        "hypotheses": [],
        "study_design": "",
        "primary_outcome": "",
        "secondary_outcomes": [],
        "funding": "",
        "conflicts": "",
        "ethics_permissions": "",
        "data_availability": "",
        "code_availability": "",
        "keywords": [],
        "selected_methods": [],
        "selected_artifacts": [],
        "figures": [],
        "tables": [],
        "abstract_background": "",
        "abstract_methods": "",
        "abstract_results": "",
        "abstract_conclusion": "",
        "introduction": "",
        "methods_notes": "",
        "results": "",
        "discussion": "",
        "limitations": "",
        "conclusion": "",
        "acknowledgements": "",
        "created_at": now,
        "updated_at": now,
    }


class StudyStore:
    """Legacy JSON store retained for backward compatibility/import into 11.15."""
    def __init__(self, root: str | Path):
        self.root = Path(root)
        self.root.mkdir(parents=True, exist_ok=True)

    def path_for(self, study_id: str) -> Path:
        return self.root / f"{study_id}.json"

    def save(self, study: Mapping[str, Any]) -> dict[str, Any]:
        payload = deepcopy(dict(study))
        payload.setdefault("study_id", str(uuid.uuid4()))
        payload.setdefault("created_at", utc_now_iso())
        payload["updated_at"] = utc_now_iso()
        target = self.path_for(payload["study_id"])
        temporary = target.with_name(f"{target.name}.{os.getpid()}.{time.time_ns()}.tmp")
        try:
            temporary.write_text(json.dumps(json_safe(payload), indent=2, ensure_ascii=False), encoding="utf-8")
            last_error = None
            for attempt in range(10):
                try:
                    os.replace(temporary, target)
                    break
                except (PermissionError, OSError) as error:
                    last_error = error
                    time.sleep(0.12 * (attempt + 1))
            else:
                raise PublicationBuilderError(f"Could not save study: {last_error}")
        finally:
            temporary.unlink(missing_ok=True)
        return payload

    def load(self, study_id: str) -> dict[str, Any]:
        path = self.path_for(study_id)
        if not path.exists():
            raise PublicationBuilderError(f"Study not found: {study_id}")
        return json.loads(path.read_text(encoding="utf-8"))

    def list_studies(self) -> pd.DataFrame:
        rows = []
        for path in sorted(self.root.glob("*.json")):
            try:
                study = json.loads(path.read_text(encoding="utf-8"))
                rows.append({"Study ID": study.get("study_id"), "Title": study.get("title"), "Project ID": study.get("project_id"), "Journal": study.get("journal"), "Updated": study.get("updated_at")})
            except Exception as error:
                rows.append({"Study ID": path.stem, "Title": f"Unreadable: {error}"})
        return pd.DataFrame(rows)


def discover_dataframes(value: Any, prefix: str = "project") -> dict[str, pd.DataFrame]:
    found: dict[str, pd.DataFrame] = {}
    frame = dataframe_from_json_safe(value)
    if frame is not None:
        found[prefix] = frame
        return found
    if isinstance(value, Mapping):
        for key, child in value.items():
            found.update(discover_dataframes(child, f"{prefix}.{key}"))
    elif isinstance(value, list):
        for index, child in enumerate(value):
            found.update(discover_dataframes(child, f"{prefix}[{index}]"))
    return found


def artifact_catalog(project: Mapping[str, Any] | None, session_artifacts: Mapping[str, Any] | None = None) -> tuple[pd.DataFrame, dict[str, pd.DataFrame]]:
    frames: dict[str, pd.DataFrame] = {}
    if project:
        frames.update(discover_dataframes(project, "project"))
    if session_artifacts:
        frames.update(discover_dataframes(session_artifacts, "session"))
    rows = [{"Artifact": name, "Rows": len(frame), "Columns": len(frame.columns), "Column names": ", ".join(map(str, frame.columns[:12])) + (" …" if len(frame.columns) > 12 else "")} for name, frame in frames.items()]
    return pd.DataFrame(rows), frames


def _numeric(frame: pd.DataFrame, column: str) -> pd.Series:
    return pd.to_numeric(frame[column], errors="coerce")


def _figure_bytes(
    frame: pd.DataFrame,
    *,
    chart_type: str,
    x_column: str,
    y_columns: Sequence[str],
    group_column: str | None = None,
    title: str = "",
    dpi: int = 300,
    lower_column: str | None = None,
    upper_column: str | None = None,
    error_column: str | None = None,
    output_format: str = "png",
) -> bytes:
    if frame.empty:
        raise PublicationBuilderError("Figure data are empty.")
    if x_column not in frame.columns and chart_type not in ("Box", "Violin", "Confusion matrix"):
        raise PublicationBuilderError("Selected x column is not present in the figure data.")
    if not y_columns and chart_type not in ("Confusion matrix",):
        raise PublicationBuilderError("Select at least one y column.")
    data = frame.copy()
    for column in y_columns:
        if column in data.columns:
            data[column] = _numeric(data, column)
    fig, ax = plt.subplots(figsize=(7.4, 4.8))
    kind = str(chart_type or "Line")
    if kind in ("Scatter", "PCA / climate space", "Spatial points"):
        y = y_columns[0]
        if group_column and group_column in data.columns:
            for group, subset in data.groupby(group_column, dropna=False):
                ax.scatter(subset[x_column], subset[y], label=str(group), alpha=0.78)
            ax.legend(fontsize=8)
        else:
            ax.scatter(data[x_column], data[y], alpha=0.78)
    elif kind == "Observed vs predicted":
        y = y_columns[0]
        x = _numeric(data, x_column)
        pred = _numeric(data, y)
        ax.scatter(x, pred, alpha=0.72)
        finite = pd.concat([x, pred], ignore_index=True).dropna()
        if not finite.empty:
            lo, hi = float(finite.min()), float(finite.max())
            ax.plot([lo, hi], [lo, hi], linestyle="--", linewidth=1)
        ax.set_ylabel(str(y))
    elif kind == "Residuals":
        pred = _numeric(data, x_column)
        observed = _numeric(data, y_columns[0])
        residual = observed - pred
        ax.scatter(pred, residual, alpha=0.72)
        ax.axhline(0, linestyle="--", linewidth=1)
        ax.set_ylabel("Observed - predicted")
    elif kind == "Bar":
        data.set_index(x_column)[list(y_columns)].plot(kind="bar", ax=ax)
    elif kind == "Box":
        data[list(y_columns)].plot(kind="box", ax=ax)
    elif kind == "Violin":
        arrays = [data[c].dropna().to_numpy() for c in y_columns]
        ax.violinplot(arrays, showmedians=True)
        ax.set_xticks(range(1, len(y_columns) + 1), labels=[str(c) for c in y_columns])
    elif kind == "Error bars":
        y = y_columns[0]
        yerr = _numeric(data, error_column) if error_column and error_column in data.columns else None
        ax.errorbar(data[x_column], data[y], yerr=yerr, marker="o", linestyle="none", capsize=3)
    elif kind == "Time series with interval":
        y = y_columns[0]
        ax.plot(data[x_column], data[y], marker="o", linewidth=1.5, label=str(y))
        if lower_column in data.columns and upper_column in data.columns:
            lower = _numeric(data, lower_column)
            upper = _numeric(data, upper_column)
            ax.fill_between(data[x_column], lower, upper, alpha=0.18)
    elif kind == "Calibration":
        observed = _numeric(data, x_column)
        predicted = _numeric(data, y_columns[0])
        tmp = pd.DataFrame({"o": observed, "p": predicted}).dropna()
        if tmp.empty:
            raise PublicationBuilderError("No complete observed/predicted pairs for calibration plot.")
        tmp["bin"] = pd.qcut(tmp["p"], q=min(10, max(2, tmp["p"].nunique())), duplicates="drop")
        cal = tmp.groupby("bin", observed=False).agg(observed=("o", "mean"), predicted=("p", "mean"))
        ax.plot(cal["predicted"], cal["observed"], marker="o")
        finite = pd.concat([cal["observed"], cal["predicted"]]).dropna()
        if not finite.empty:
            lo, hi = float(finite.min()), float(finite.max())
            ax.plot([lo, hi], [lo, hi], linestyle="--", linewidth=1)
        ax.set_xlabel("Mean predicted")
        ax.set_ylabel("Mean observed")
    elif kind in ("ROC", "Precision-recall"):
        truth = _numeric(data, x_column)
        score = _numeric(data, y_columns[0])
        tmp = pd.DataFrame({"truth": truth, "score": score}).dropna()
        if not set(tmp["truth"].unique()).issubset({0, 1}) or tmp["truth"].nunique() != 2:
            raise PublicationBuilderError(f"{kind} requires binary 0/1 observations and probability/scores.")
        try:
            from sklearn.metrics import precision_recall_curve, roc_curve
        except Exception as exc:
            raise PublicationBuilderError("scikit-learn is required for ROC/PR figures.") from exc
        if kind == "ROC":
            fpr, tpr, _ = roc_curve(tmp["truth"], tmp["score"])
            ax.plot(fpr, tpr)
            ax.plot([0, 1], [0, 1], linestyle="--", linewidth=1)
            ax.set_xlabel("False-positive rate")
            ax.set_ylabel("True-positive rate")
        else:
            precision, recall, _ = precision_recall_curve(tmp["truth"], tmp["score"])
            ax.plot(recall, precision)
            ax.set_xlabel("Recall")
            ax.set_ylabel("Precision")
    elif kind == "Confusion matrix":
        if len(y_columns) < 1:
            raise PublicationBuilderError("Confusion matrix requires observed and predicted label columns.")
        observed = data[x_column].astype(str)
        predicted = data[y_columns[0]].astype(str)
        labels = sorted(set(observed.dropna()) | set(predicted.dropna()))
        matrix = pd.crosstab(observed, predicted).reindex(index=labels, columns=labels, fill_value=0)
        image = ax.imshow(matrix.to_numpy(), aspect="auto")
        ax.set_xticks(range(len(labels)), labels=labels, rotation=45, ha="right")
        ax.set_yticks(range(len(labels)), labels=labels)
        ax.set_xlabel("Predicted")
        ax.set_ylabel("Observed")
        for i in range(len(labels)):
            for j in range(len(labels)):
                ax.text(j, i, str(matrix.iloc[i, j]), ha="center", va="center")
        fig.colorbar(image, ax=ax, fraction=0.046, pad=0.04)
    else:
        for column in y_columns:
            ax.plot(data[x_column], data[column], marker="o", linewidth=1.5, label=str(column))
        if len(y_columns) > 1:
            ax.legend(fontsize=8)
    ax.set_title(title or kind)
    if kind not in ("Calibration", "ROC", "Precision-recall", "Confusion matrix"):
        ax.set_xlabel(str(x_column))
        if kind not in ("Residuals", "Observed vs predicted"):
            ax.set_ylabel(", ".join(map(str, y_columns)))
    ax.grid(alpha=0.18)
    fig.tight_layout()
    buffer = io.BytesIO()
    fig.savefig(buffer, format=output_format, dpi=int(dpi), bbox_inches="tight")
    plt.close(fig)
    return buffer.getvalue()


def figure_png(frame: pd.DataFrame, *, chart_type: str, x_column: str, y_columns: Sequence[str], group_column: str | None = None, title: str = "", dpi: int = 300, lower_column: str | None = None, upper_column: str | None = None, error_column: str | None = None) -> bytes:
    return _figure_bytes(frame, chart_type=chart_type, x_column=x_column, y_columns=y_columns, group_column=group_column, title=title, dpi=dpi, lower_column=lower_column, upper_column=upper_column, error_column=error_column, output_format="png")


def figure_svg(frame: pd.DataFrame, *, chart_type: str, x_column: str, y_columns: Sequence[str], group_column: str | None = None, title: str = "", lower_column: str | None = None, upper_column: str | None = None, error_column: str | None = None) -> bytes:
    return _figure_bytes(frame, chart_type=chart_type, x_column=x_column, y_columns=y_columns, group_column=group_column, title=title, dpi=300, lower_column=lower_column, upper_column=upper_column, error_column=error_column, output_format="svg")


def multi_panel_png(figures: Sequence[bytes], labels: Sequence[str] | None = None, dpi: int = 300) -> bytes:
    if not figures:
        raise PublicationBuilderError("No figure panels were supplied.")
    panels = list(figures)[:6]
    count = len(panels)
    cols = 2 if count > 1 else 1
    rows = int(np.ceil(count / cols))
    fig, axes = plt.subplots(rows, cols, figsize=(7.4 * cols / 1.35, 4.6 * rows))
    axes_array = np.atleast_1d(axes).reshape(-1)
    for idx, (axis, raw) in enumerate(zip(axes_array, panels)):
        image = plt.imread(io.BytesIO(raw), format="png")
        axis.imshow(image)
        axis.axis("off")
        if labels and idx < len(labels):
            axis.set_title(str(labels[idx]), loc="left", fontweight="bold")
    for axis in axes_array[len(panels):]:
        axis.axis("off")
    fig.tight_layout()
    out = io.BytesIO()
    fig.savefig(out, format="png", dpi=int(dpi), bbox_inches="tight")
    plt.close(fig)
    return out.getvalue()


def _markdown_table(frame: pd.DataFrame, max_rows: int = 30, max_columns: int = 14) -> str:
    table = frame.head(max_rows).iloc[:, :max_columns].copy()
    if table.empty:
        return "_No rows available._"
    headers = [str(c).replace("|", "\\|") for c in table.columns]
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in table.itertuples(index=False, name=None):
        values = [str(v).replace("|", "\\|") if pd.notna(v) else "" for v in row]
        lines.append("| " + " | ".join(values) + " |")
    return "\n".join(lines)


def manuscript_word_counts(study: Mapping[str, Any]) -> dict[str, int]:
    keys = ["abstract_background", "abstract_methods", "abstract_results", "abstract_conclusion", "introduction", "methods_notes", "results", "discussion", "limitations", "conclusion"]
    counts = {key: len(re.findall(r"\b\w+[\w'-]*\b", str(study.get(key) or ""))) for key in keys}
    counts["abstract_total"] = sum(counts[k] for k in keys[:4])
    counts["main_text"] = sum(counts[k] for k in keys[4:])
    return counts


HIGH_RISK_CLAIM_TERMS = {
    "caused": "Causal wording requires a causal design/estimate rather than predictive association alone.",
    "causes": "Causal wording requires a causal design/estimate rather than predictive association alone.",
    "improved": "State the comparator and effect estimate; predictive performance alone does not establish an agronomic treatment effect.",
    "optimal": "An optimum requires a validated objective, constraints and applicability range.",
    "validated": "Specify the validation scope (within trial, cross-season, cross-site, external, etc.).",
    "significant": "Link the statistical model/test, estimate, uncertainty and multiplicity handling where relevant.",
    "generalizable": "Specify the independent environments/populations used to assess generalisation.",
    "guaranteed": "Agronomic predictions and timing synchrony do not guarantee biological or commercial outcomes.",
}


def audit_claim_text(text: str, *, evidence_type: str = "", statistic: str = "", source_reference: str = "") -> list[str]:
    lower = str(text or "").casefold()
    warnings = []
    for term, message in HIGH_RISK_CLAIM_TERMS.items():
        if re.search(rf"\b{re.escape(term)}\b", lower):
            if term in ("caused", "causes") and "causal" in str(evidence_type).casefold():
                continue
            warnings.append(message)
    if text and not source_reference.strip():
        warnings.append("No evidence source is linked to this claim.")
    if "significant" in lower and not statistic.strip():
        warnings.append("The claim uses 'significant' but no effect/statistic is recorded in the claim ledger.")
    return list(dict.fromkeys(warnings))


def report_audit(study: Mapping[str, Any], *, claims: Sequence[Mapping[str, Any]] | None = None, figure_count: int = 0, table_count: int = 0, citation_count: int = 0, snapshot_present: bool = False, evidence_manifest: Mapping[str, Any] | None = None) -> list[dict[str, str]]:
    issues: list[dict[str, str]] = []
    manuscript_fields = ["title", "abstract_background", "abstract_methods", "abstract_results", "abstract_conclusion", "introduction", "discussion", "limitations", "conclusion"]
    combined = "\n".join(str(study.get(k) or "") for k in manuscript_fields)
    placeholders = sorted(set(re.findall(r"\[\[[^\]]+\]\]", combined)))
    for value in placeholders:
        issues.append({"severity": "Warning", "category": "Placeholder", "message": f"Unresolved manuscript placeholder: {value}"})
    if not snapshot_present:
        issues.append({"severity": "Warning", "category": "Reproducibility", "message": "No frozen evidence snapshot is linked to the current report version."})
    if not str(study.get("limitations") or "").strip():
        issues.append({"severity": "Warning", "category": "Scientific completeness", "message": "Limitations section is empty."})
    if not str(study.get("data_availability") or "").strip():
        issues.append({"severity": "Info", "category": "Reproducibility", "message": "Data-availability statement is empty."})
    if citation_count == 0:
        issues.append({"severity": "Info", "category": "Citations", "message": "No citations are linked to the report."})
    for claim in claims or []:
        for warning in audit_claim_text(str(claim.get("claim_text") or claim.get("text") or ""), evidence_type=str(claim.get("evidence_type") or ""), statistic=str(claim.get("statistic") or ""), source_reference=str(claim.get("source_reference") or "")):
            issues.append({"severity": "Warning", "category": "Claim", "message": f"{str(claim.get('claim_text') or claim.get('text') or '')[:90]} — {warning}"})
    if evidence_manifest:
        counts = evidence_manifest.get("counts") or {}
        if int(counts.get("training_runs", 0) or 0) > 1:
            issues.append({"severity": "Info", "category": "Selective reporting", "message": "Multiple model training runs are present in the frozen evidence. Report the model-selection process rather than only the winning model."})
        readiness = evidence_manifest.get("readiness") or {}
        for name, status in readiness.items():
            if isinstance(status, Mapping) and status.get("status") in ("Missing", "Incomplete"):
                issues.append({"severity": "Warning", "category": "Evidence readiness", "message": f"{name}: {status.get('message') or status.get('status')}"})
    if figure_count == 0 and table_count == 0:
        issues.append({"severity": "Info", "category": "Reporting", "message": "No registered tables or figures are linked to the report."})
    return issues


def redact_text(text: str, options: Mapping[str, bool] | None = None) -> str:
    options = dict(options or {})
    result = str(text or "")
    if options.get("coordinates"):
        result = re.sub(r"(?<!\d)([-+]?\d{1,2}\.\d{3,})\s*[,;/ ]\s*([-+]?\d{1,3}\.\d{3,})(?!\d)", "[[REDACTED COORDINATES]]", result)
    return result


def redact_frame(frame: pd.DataFrame, options: Mapping[str, bool] | None = None) -> tuple[pd.DataFrame, list[str]]:
    options = dict(options or {})
    out = frame.copy()
    removed: list[str] = []
    for column in list(out.columns):
        norm = re.sub(r"[^a-z0-9]+", "_", str(column).casefold()).strip("_")
        should = False
        if options.get("coordinates") and any(token in norm for token in ("latitude", "longitude", "lat", "lon", "geometry", "geojson")):
            should = True
        if options.get("field_names") and any(token in norm for token in ("field_name", "farm_name", "research_centre", "research_center")):
            should = True
        if options.get("genotypes") and any(token in norm for token in ("genotype", "parent", "variety", "cultivar", "line_name")):
            should = True
        if options.get("researcher_names") and any(token in norm for token in ("observer", "operator", "author", "researcher", "user_name", "email")):
            should = True
        if should:
            out[column] = "[[REDACTED]]"
            removed.append(str(column))
    return out, removed


def manuscript_markdown(*, study: Mapping[str, Any], selected_methods: Sequence[str], tables: Sequence[Mapping[str, Any]], figures: Sequence[Mapping[str, Any]], reproducibility: Mapping[str, Any], citations: Sequence[Mapping[str, Any]] | None = None, claims: Sequence[Mapping[str, Any]] | None = None) -> str:
    authors_raw = study.get("authors", [])
    if authors_raw and isinstance(authors_raw[0], Mapping):
        authors = ", ".join(str(a.get("name") or "") for a in authors_raw if a.get("name"))
    else:
        authors = ", ".join(map(str, authors_raw))
    affiliations = study.get("affiliations") or []
    lines = [
        f"# {study.get('title', 'Untitled study')}", "",
        f"**Authors:** {authors or '[[AUTHORS]]'}",
        f"**Affiliations:** {'; '.join(map(str, affiliations)) if affiliations else '[[AFFILIATIONS]]'}",
        f"**Corresponding author:** {study.get('corresponding_author') or '[[CORRESPONDING AUTHOR]]'}",
        f"**Target journal:** {study.get('journal') or study.get('target_journal') or '[[JOURNAL]]'}",
        f"**Report type:** {study.get('report_type') or 'Scientific report'}", "",
        "## Abstract", "",
        f"**Background:** {study.get('abstract_background') or '[[Summarise the scientific context.]]'}", "",
        f"**Methods:** {study.get('abstract_methods') or '[[Summarise data, models and validation design.]]'}", "",
        f"**Results:** {study.get('abstract_results') or '[[Insert the main numerical findings after final analyses.]]'}", "",
        f"**Conclusions:** {study.get('abstract_conclusion') or '[[State the supported conclusion without exceeding the evidence.]]'}", "",
    ]
    keywords = study.get("keywords") or []
    if keywords:
        lines.extend([f"**Keywords:** {', '.join(map(str, keywords))}", ""])
    lines.extend(["## 1. Introduction", "", study.get("introduction") or "[[Insert the introduction.]]", "", "## 2. Materials and Methods", ""])
    if study.get("study_design"):
        lines.extend(["### 2.1 Study design", "", str(study.get("study_design")), ""])
    if study.get("research_question"):
        lines.extend(["### 2.2 Research question", "", str(study.get("research_question")), ""])
    method_start = 3
    for index, method in enumerate(selected_methods, start=method_start):
        record = METHODS_REGISTRY.get(method, {"text": "[[Method text required.]]", "references": [], "relationship": "Unregistered method"})
        lines.extend([f"### 2.{index} {method}", "", record["text"], "", f"*Implementation relationship:* {record.get('relationship','')}", ""])
    if study.get("methods_notes"):
        lines.extend(["### Additional method details", "", str(study.get("methods_notes")), ""])
    lines.extend(["## 3. Results", "", study.get("results") or "Results are linked to registered tables, figures and the claim ledger; authors must verify every numerical statement before submission.", ""])
    for index, table in enumerate(tables, start=1):
        lines.extend([f"### 3.{index} {table.get('title', f'Table {index}')}", "", table.get("narrative", ""), "", f"**Table {index}. {table.get('caption', '')}**", "", table.get("markdown", ""), ""])
    lines.extend(["## 4. Discussion", "", study.get("discussion") or "[[Interpret the findings, compare with prior work, and distinguish evidence from inference.]]", "", "## 5. Limitations", "", study.get("limitations") or "[[State limitations supported by the evidence and model cards.]]", "", "## 6. Conclusions", "", study.get("conclusion") or "[[Insert the conclusion supported by the final evidence.]]", ""])
    for heading, key, fallback in [
        ("Acknowledgements", "acknowledgements", ""), ("Funding", "funding", "[[Funding statement]]"), ("Conflicts of interest", "conflicts", "[[Conflict-of-interest statement]]"), ("Ethics and permissions", "ethics_permissions", "[[Ethics/permissions statement if applicable]]"), ("Data availability", "data_availability", "[[Describe access, licences and restrictions.]]"), ("Code availability", "code_availability", "[[Describe code/model package access and version.]]")
    ]:
        lines.extend([f"## {heading}", "", study.get(key) or fallback, ""])
    if claims:
        lines.extend(["## Claim-evidence register", ""])
        for idx, claim in enumerate(claims, 1):
            lines.append(f"- **Claim {idx}:** {claim.get('claim_text') or claim.get('text')} — Evidence: {claim.get('source_reference') or 'not linked'}; {claim.get('statistic') or 'statistic not recorded'}.")
        lines.append("")
    lines.extend(["## Reproducibility manifest", "", "```json", json.dumps(json_safe(dict(reproducibility)), indent=2, ensure_ascii=False), "```", "", "## References", ""])
    references: list[str] = []
    for method in selected_methods:
        references.extend(METHODS_REGISTRY.get(method, {}).get("references", []))
    for citation in citations or []:
        authors_text = citation.get("authors") or ""
        year = citation.get("year") or ""
        title = citation.get("title") or ""
        journal = citation.get("journal") or ""
        doi = citation.get("doi") or ""
        references.append(f"{authors_text} ({year}). {title}. {journal}." + (f" DOI {doi}." if doi else ""))
    for ref in dict.fromkeys(r for r in references if r):
        lines.append(f"- {ref}")
    lines.extend(["", "## Figure register", ""])
    for index, figure in enumerate(figures, start=1):
        lines.append(f"- **Figure {index}.** {figure.get('caption', '') or figure.get('title','')}")
    return "\n".join(lines)


def manuscript_html(markdown_text: str) -> str:
    escaped = html.escape(markdown_text)
    body = "<pre>" + escaped + "</pre>"
    return "<!doctype html><html><head><meta charset='utf-8'><title>AGROLATTICE report</title><style>body{font-family:Arial,sans-serif;max-width:1000px;margin:40px auto;line-height:1.5}pre{white-space:pre-wrap;font-family:Arial,sans-serif}</style></head><body>" + body + "</body></html>"


def manuscript_docx(*, study: Mapping[str, Any], markdown_text: str, tables: Sequence[Mapping[str, Any]], figures: Sequence[Mapping[str, Any]], selected_methods: Sequence[str] | None = None, citations: Sequence[Mapping[str, Any]] | None = None) -> bytes:
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except Exception as error:
        raise PublicationBuilderError("DOCX export requires python-docx.") from error
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    title = document.add_heading(study.get("title", "Untitled study"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors_raw = study.get("authors", [])
    if authors_raw and isinstance(authors_raw[0], Mapping):
        authors = ", ".join(str(a.get("name") or "") for a in authors_raw if a.get("name"))
    else:
        authors = ", ".join(map(str, authors_raw))
    p = document.add_paragraph(authors or "[[AUTHORS]]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff = study.get("affiliations") or []
    if aff:
        pa = document.add_paragraph("; ".join(map(str, aff)))
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Target journal: {study.get('journal') or study.get('target_journal') or '[[JOURNAL]]'}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_heading("Abstract", level=1)
    for label, key in [("Background", "abstract_background"), ("Methods", "abstract_methods"), ("Results", "abstract_results"), ("Conclusions", "abstract_conclusion")]:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(study.get(key) or "[[To be completed.]]")
    document.add_heading("1. Introduction", level=1)
    document.add_paragraph(study.get("introduction") or "[[Introduction required.]]")
    document.add_heading("2. Materials and Methods", level=1)
    if study.get("study_design"):
        document.add_heading("2.1 Study design", level=2)
        document.add_paragraph(str(study.get("study_design")))
    for index, method in enumerate(selected_methods or study.get("selected_methods", []), start=2):
        record = METHODS_REGISTRY.get(method, {})
        document.add_heading(f"2.{index} {method}", level=2)
        document.add_paragraph(record.get("text") or "[[Method text required.]]")
        rel = record.get("relationship")
        if rel:
            pr = document.add_paragraph()
            pr.add_run("Implementation relationship: ").italic = True
            pr.add_run(str(rel)).italic = True
    document.add_heading("3. Results", level=1)
    document.add_paragraph(study.get("results") or "Results are linked to registered evidence and require author verification.")
    for index, item in enumerate(tables, start=1):
        document.add_heading(item.get("title", f"Result table {index}"), level=2)
        if item.get("narrative"):
            document.add_paragraph(str(item.get("narrative")))
        document.add_paragraph(f"Table {index}. {item.get('caption', '')}", style="Caption")
        frame = item.get("frame")
        if isinstance(frame, pd.DataFrame) and not frame.empty:
            display = frame.head(35).iloc[:, :12]
            table = document.add_table(rows=1, cols=len(display.columns))
            table.style = "Table Grid"
            for cell, column in zip(table.rows[0].cells, display.columns):
                cell.text = str(column)
            for row in display.itertuples(index=False, name=None):
                cells = table.add_row().cells
                for cell, value in zip(cells, row):
                    cell.text = "" if pd.isna(value) else str(value)
    for index, item in enumerate(figures, start=1):
        image = item.get("png")
        if image:
            document.add_picture(io.BytesIO(image), width=Inches(6.4))
            document.add_paragraph(f"Figure {index}. {item.get('caption', '') or item.get('title','')}", style="Caption")
    for idx, (heading, key, fallback) in enumerate([
        ("4. Discussion", "discussion", "[[Discussion required.]]"),
        ("5. Limitations", "limitations", "[[Limitations required.]]"),
        ("6. Conclusions", "conclusion", "[[Conclusion required.]]"),
        ("Acknowledgements", "acknowledgements", ""),
        ("Funding", "funding", "[[Funding statement]]"),
        ("Conflicts of interest", "conflicts", "[[Conflict-of-interest statement]]"),
        ("Data availability", "data_availability", "[[Data availability statement]]"),
        ("Code availability", "code_availability", "[[Code availability statement]]"),
    ]):
        document.add_heading(heading, level=1)
        document.add_paragraph(study.get(key) or fallback)
    document.add_heading("References", level=1)
    reference_lines: list[str] = []
    for method in selected_methods or study.get("selected_methods", []):
        reference_lines.extend(METHODS_REGISTRY.get(method, {}).get("references", []))
    for citation in citations or []:
        ref = f"{citation.get('authors','')} ({citation.get('year','')}). {citation.get('title','')}. {citation.get('journal','')}."
        if citation.get("doi"):
            ref += f" DOI {citation.get('doi')}."
        reference_lines.append(ref)
    for ref in dict.fromkeys(r for r in reference_lines if r):
        document.add_paragraph(ref)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_publication_package(*, study: Mapping[str, Any], selected_methods: Sequence[str], selected_tables: Sequence[Mapping[str, Any]], figures: Sequence[Mapping[str, Any]], reproducibility: Mapping[str, Any], project: Mapping[str, Any] | None = None, citations: Sequence[Mapping[str, Any]] | None = None, claims: Sequence[Mapping[str, Any]] | None = None, privacy_profile: str = "Internal research package", redaction_options: Mapping[str, bool] | None = None, package_type: str = "Complete research package") -> bytes:
    table_records = []
    redaction_report: dict[str, Any] = {"privacy_profile": privacy_profile, "redacted_columns": {}}
    for item in selected_tables:
        frame = item.get("frame")
        effective = frame
        if isinstance(frame, pd.DataFrame) and privacy_profile.startswith("Public"):
            effective, removed = redact_frame(frame, redaction_options)
            if removed:
                redaction_report["redacted_columns"][str(item.get("title") or "table")] = removed
        table_records.append({**dict(item), "frame": effective, "markdown": _markdown_table(effective) if isinstance(effective, pd.DataFrame) else ""})
    study_payload = deepcopy(dict(study))
    study_payload["selected_methods"] = list(selected_methods)
    if privacy_profile.startswith("Public"):
        if (redaction_options or {}).get("researcher_names"):
            study_payload["authors"] = ["[[REDACTED FOR PUBLIC PACKAGE]]"]
            study_payload["corresponding_author"] = "[[REDACTED FOR PUBLIC PACKAGE]]"
        for key in ("introduction", "methods_notes", "results", "discussion", "limitations", "conclusion"):
            study_payload[key] = redact_text(str(study_payload.get(key) or ""), redaction_options)
    markdown_text = manuscript_markdown(study=study_payload, selected_methods=selected_methods, tables=table_records, figures=figures, reproducibility=reproducibility, citations=citations, claims=claims)
    html_text = manuscript_html(markdown_text)
    docx_bytes = manuscript_docx(study=study_payload, markdown_text=markdown_text, tables=table_records, figures=figures, selected_methods=selected_methods, citations=citations)
    audit_rows = report_audit(study_payload, claims=claims, figure_count=len(figures), table_count=len(table_records), citation_count=len(citations or []), snapshot_present=bool(reproducibility.get("evidence_snapshot_id")), evidence_manifest=reproducibility)
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("manuscript/manuscript.md", markdown_text)
        archive.writestr("manuscript/manuscript.html", html_text)
        archive.writestr("manuscript/manuscript.docx", docx_bytes)
        archive.writestr("study_protocol.json", json.dumps(json_safe(study_payload), indent=2, ensure_ascii=False))
        archive.writestr("reproducibility/reproducibility_manifest.json", json.dumps(json_safe(reproducibility), indent=2, ensure_ascii=False))
        archive.writestr("reproducibility/report_audit.csv", pd.DataFrame(audit_rows).to_csv(index=False))
        archive.writestr("reproducibility/redaction_report.json", json.dumps(redaction_report, indent=2, ensure_ascii=False))
        if claims:
            archive.writestr("evidence/claim_ledger.csv", pd.DataFrame(list(claims)).to_csv(index=False))
        if citations:
            archive.writestr("references/citations.csv", pd.DataFrame(list(citations)).to_csv(index=False))
            bibtex = "\n\n".join(str(c.get("bibtex") or "") for c in citations if c.get("bibtex"))
            ris = "\n\n".join(str(c.get("ris") or "") for c in citations if c.get("ris"))
            if bibtex:
                archive.writestr("references/references.bib", bibtex)
            if ris:
                archive.writestr("references/references.ris", ris)
        if project:
            archive.writestr("legacy/project_snapshot.json", json.dumps(json_safe(project), indent=2, ensure_ascii=False))
        for index, item in enumerate(table_records, start=1):
            frame = item.get("frame")
            if isinstance(frame, pd.DataFrame):
                archive.writestr(f"tables/table_{index:02d}_{slugify(item.get('title', 'table'))}.csv", frame.to_csv(index=False))
                archive.writestr(f"tables/table_{index:02d}_{slugify(item.get('title', 'table'))}_manifest.json", json.dumps(json_safe({"source": item.get("source"), "settings": item.get("settings"), "caption": item.get("caption")}), indent=2))
        for index, item in enumerate(figures, start=1):
            if item.get("png"):
                archive.writestr(f"figures/figure_{index:02d}_{slugify(item.get('title', 'figure'))}.png", item["png"])
            if item.get("svg"):
                archive.writestr(f"figures/figure_{index:02d}_{slugify(item.get('title', 'figure'))}.svg", item["svg"])
            archive.writestr(f"figures/figure_{index:02d}_{slugify(item.get('title', 'figure'))}_manifest.json", json.dumps(json_safe({k: v for k, v in item.items() if k not in ("png", "svg")}), indent=2))
        method_rows = []
        for method in selected_methods:
            record = METHODS_REGISTRY.get(method, {})
            method_rows.append({"Method": method, "Implementation relationship": record.get("relationship"), "Text": record.get("text"), "References": " | ".join(record.get("references", []))})
        archive.writestr("methods/methods_provenance.csv", pd.DataFrame(method_rows).to_csv(index=False))
        readme = (
            f"AGROLATTICE research reporting package\nPackage type: {package_type}\nPrivacy profile: {privacy_profile}\n\n"
            "All narrative text, numerical claims, statistical interpretations, citations and journal-specific formatting require author review before submission.\n"
            "Measured, derived, mechanistic, predictive, scenario, recommendation and causal-estimate evidence must not be treated as interchangeable.\n"
        )
        archive.writestr("README.txt", readme)
    return buffer.getvalue()
